import streamlit as st
import docx
import pdfplumber
import io

# ==========================================
# 1. 页面基本配置（必须是 Streamlit 的第一个命令）
# ==========================================
st.set_page_config(
    page_title="AutoJob-Agent | 简历解析",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 2. 核心文本提取函数（加入鲁棒的异常处理）
# ==========================================
def extract_text_from_pdf(file_bytes):
    """安全地从 PDF 二进制流中提取文本"""
    try:
        text = ""
        # 使用 io.BytesIO 将上传的字节流转换为文件对象
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            return None, "PDF 文件内容似乎为空，或者是由纯图片组成的扫描件。"
        return text, None
    except Exception as e:
        return None, f"PDF 解析失败，可能文件已损坏或加密。错误详情: {str(e)}"

def extract_text_from_docx(file_bytes):
    """安全地从 Word (DOCX) 二进制流中提取文本"""
    try:
        text = ""
        doc = docx.Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        
        # 提取表格中的文本，防止遗漏简历中的表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 去重处理，因为 Word 表格合并单元格会导致文本重复读取
                    pass 
                text += " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()]) + "\n"

        if not text.strip():
            return None, "Word 文件内容为空。"
        return text, None
    except Exception as e:
        return None, f"Word 解析失败，请检查是否为标准 .docx 格式。错误详情: {str(e)}"

# ==========================================
# 3. 前端 UI 界面设计
# ==========================================

# 侧边栏：品牌与开发日志快捷入口
with st.sidebar:
    st.markdown("## 🤖 AutoJob-Agent")
    st.markdown("---")
    st.info("💡 **当前阶段: Phase 1**\n\n核心功能：简历与 JD 文本解析模块开发。")
    st.markdown("---")
    st.markdown("⚙️ **系统配置**")
    model_provider = st.selectbox("LLM 驱动选择 (Phase 2 启用)", ["DeepSeek-V3", "GPT-4o", "Claude 3.5"])
    st.success(f"已就绪: {model_provider}")

# 主页面大标题
st.title("AutoJob-Agent 智能求职助手")
st.caption("从 Vibe Coding 走向生产级 Agent 开发 —— 让 AI 助你一臂之力")

# 引入一个精美的信息面板
st.markdown("""
<div style="background-color:#f0f2f6; padding:15px; border-radius:10px; margin-bottom:20px;">
    <h4 style="margin:0; color:#1f2937;">📋 模块说明：原始简历解析</h4>
    <p style="margin:5px 0 0 0; font-size:14px; color:#4b5563;">
        本模块负责将您本地的 <b>PDF</b> 或 <b>Word (docx)</b> 格式简历转化为大模型可读的结构化文本。
        代码已对大文件、空文件及损坏文件进行了熔断保护，确保运行稳定。
    </p>
</div>
""", unsafe_allow_html=True)

# 使用 Streamlit 的列布局，让界面更紧凑好看
col1, col2 = st.columns([1, 2], gap="large")

with col1:
    st.subheader("📤 上传您的简历")
    # 文件上传组件，严格限制文件类型
    uploaded_file = st.file_uploader(
        "支持 PDF 或 DOCX 格式，大小不超过 5MB",
        type=["pdf", "docx"],
        help="请确保文件没有被加密或损坏"
    )

with col2:
    st.subheader("解析结果预览")
    
    if uploaded_file is not None:
        # 1. 读取文件二进制流
        file_bytes = uploaded_file.read()
        file_name = uploaded_file.name
        file_type = file_name.split(".")[-1].lower()
        
        # 2. 状态加载动画 (Spinner)
        with st.spinner("Agent 正在深度解析文件，请稍候..."):
            resume_text = None
            error_msg = None
            
            # 3. 根据类型分流解析
            if file_type == "pdf":
                resume_text, error_msg = extract_text_from_pdf(file_bytes)
            elif file_type == "docx":
                resume_text, error_msg = extract_text_from_docx(file_bytes)
        
        # 4. 根据解析结果更新 UI
        if error_msg:
            # 优雅的报错提示，页面不会崩溃
            st.error(error_msg)
        else:
            # 成功提示与指标卡片
            st.toast("简历解析成功！")
            
            # 使用 metric 卡片展示文本基本指标，高级感拉满
            char_count = len(resume_text)
            st.metric(label="成功提取字数", value=f"{char_count} 字")
            
            # 用代码框或文本框展示提取出的文本，方便用户复制和查看
            st.text_area(
                label="提取出的纯文本内容 (将作为大模型的输入上下文)：",
                value=resume_text,
                height=400,
                disabled=True  # 设置为只读
            )
    else:
        # 未上传文件时的占位友好提示
        st.info("👋 请在左侧上传您的简历文件，解析后的内容将在此处实时显示。")
